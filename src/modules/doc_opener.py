def isTag(text):
    return True if (('[' in text) and (']' in text)) else False


def isCloseTag(text):
    return True if ((isTag(text) and '/' in text)) else False


def stripTag(text):
    t = text.strip(']')
    t = t.strip('[')
    return t


def readDocument(name):
    from docx import Document
    #eu estive aqui
    document = Document(name)
    text = {}
    last_tag = ''
    for paragraph in document.paragraphs:
        if isCloseTag(paragraph.text):
            continue
        elif isTag(paragraph.text):
            last_tag = stripTag(paragraph.text)
            text[last_tag] = []
        else:
            text[last_tag].append(paragraph.text)
    return text
